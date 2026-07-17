from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "vermittlungs-und-kooperationsvertrag-montagepartner.md"
OUT = ROOT / "Vermittlungs-und-Kooperationsvertrag-Montagepartner.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_run(run, bold=False, italic=False, size=None, color=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_mixed_paragraph(doc, text, style=None):
    para = doc.add_paragraph(style=style)
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if idx % 2 == 1:
            run.bold = True
    return para


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title_block(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Vermittlungs- und Kooperationsvertrag")
    style_run(run, bold=True, size=22, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("für selbständige Montagepartner von Vanventure")
    style_run(run, size=12, color="555555")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    run = note.add_run(
        "Arbeitsentwurf zur anwaltlichen Prüfung. Keine Rechtsberatung."
    )
    style_run(run, italic=True, color="9B1C1C")

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    widths = [2200, 7040]
    rows = [
        ("Dokument", "Vermittlungs- und Kooperationsvertrag"),
        ("Stand", "19. Juni 2026"),
        ("Gebühren", "50 % netto der mit dem Kunden abgerechneten Montageleistung"),
    ]
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_idx])
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].add_run(rows[row_idx][col_idx])
            if col_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                cell.paragraphs[0].runs[0].bold = True


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.text = "Vanventure | Vermittlungs- und Kooperationsvertrag"
    footer.alignment = 1
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor.from_string("666666")


def add_signature_table(doc):
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    intro.add_run(
        "Die Parteien sind damit einverstanden, diesen Vertrag elektronisch zu unterzeichnen. "
        "Eine elektronische Signatur über ein geeignetes E-Signatur-Tool oder eine unterschriebene "
        "und als Scan/PDF übermittelte Fassung genügt der vereinbarten Textform, soweit gesetzlich zulässig."
    )

    hint = doc.add_paragraph()
    hint.paragraph_format.space_after = Pt(10)
    hint.add_run("Für die elektronische Unterzeichnung sollen die nachfolgenden Felder verwendet werden.")

    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    widths = [2300, 3530, 3530]
    headers = ["Feld", "Vanventure", "Montagepartner"]
    rows = [
        ("Name des Unterzeichners", "[Name ergänzen]", "[Name ergänzen]"),
        ("Firma / Rolle", "Vanventure / [Rolle ergänzen]", "[Firma / Rolle ergänzen]"),
        ("E-Mail des Unterzeichners", "[E-Mail ergänzen]", "[E-Mail ergänzen]"),
        ("Ort / Datum", "[Ort, Datum ergänzen]", "[Ort, Datum ergänzen]"),
        ("Elektronische Signatur", "[Signaturfeld Vanventure]", "[Signaturfeld Montagepartner]"),
    ]

    for col_idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[col_idx])
        set_cell_shading(cell, "F2F4F7")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(headers[col_idx])
        run.bold = True

    for row_idx, values in enumerate(rows, start=1):
        for col_idx, value in enumerate(values):
            cell = table.rows[row_idx].cells[col_idx]
            set_cell_width(cell, widths[col_idx])
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            if col_idx == 0:
                run.bold = True
            if row_idx == 5 and col_idx > 0:
                set_cell_shading(cell, "FBFBF8")
                paragraph.paragraph_format.space_before = Pt(18)
                paragraph.paragraph_format.space_after = Pt(18)


def build_docx():
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_list = False

    for raw in lines:
        line = raw.rstrip()
        if not line:
            in_list = False
            continue

        if line.startswith("# "):
            if "Entwurf:" in line:
                continue
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            doc.add_paragraph(heading, style="Heading 1")
            if heading.startswith("16. "):
                add_signature_table(doc)
                break
            continue

        if line.startswith("- "):
            para = add_mixed_paragraph(doc, line[2:].strip(), style="List Bullet")
            para.paragraph_format.space_after = Pt(4)
            in_list = True
            continue

        if line[0:2].isdigit() and ". " in line[:5]:
            para = add_mixed_paragraph(doc, line.split(". ", 1)[1].strip(), style="List Number")
            para.paragraph_format.space_after = Pt(4)
            continue

        para = add_mixed_paragraph(doc, line)
        para.paragraph_format.space_after = Pt(6)

    add_footer(doc.sections[0])
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
